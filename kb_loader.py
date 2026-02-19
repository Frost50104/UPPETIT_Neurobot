"""
Loads and parses ``Information bank.docx``.

Returns a list of DocSection objects, each corresponding to one heading block.
Inline images are extracted and saved to *images_dir* with names that include
a section/paragraph index so they can later be linked back to chunks.
"""
from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


@dataclass
class DocSection:
    """One logical section: a heading + all body text until the next heading."""

    heading: str
    heading_level: int        # 0 = no heading, 1 = Heading 1, …
    text: str                 # concatenated paragraph text
    images: list[str] = field(default_factory=list)  # absolute paths to images
    position: int = 0         # paragraph index of the heading in the document


# ─────────────────────────────────────────────────────────────────────────────
# Public API
# ─────────────────────────────────────────────────────────────────────────────

def load_knowledge_base(docx_path: str, images_dir: str) -> list[DocSection]:
    """
    Parse *docx_path* and return a list of DocSection objects.

    Raises:
        FileNotFoundError: if the file does not exist.
        ValueError:        if the file is empty or unreadable.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"Knowledge base file not found: {docx_path}")
    if path.stat().st_size == 0:
        raise ValueError(f"Knowledge base file is empty: {docx_path}")

    logger.info("Loading knowledge base from '%s'", docx_path)
    Path(images_dir).mkdir(parents=True, exist_ok=True)

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Cannot open '{docx_path}': {exc}") from exc

    sections = _parse_paragraphs(doc, images_dir)
    table_sections = _parse_tables(doc, start_position=len(doc.paragraphs))
    sections.extend(table_sections)

    if not sections:
        raise ValueError(f"No readable content found in '{docx_path}'")

    logger.info("Loaded %d sections from knowledge base", len(sections))
    return sections


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────

def _parse_paragraphs(doc: Document, images_dir: str) -> list[DocSection]:
    sections: list[DocSection] = []

    current_heading = "Общая информация"
    current_level = 0
    current_text_parts: list[str] = []
    current_images: list[str] = []

    for position, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        level = _heading_level(style_name)
        para_images = _extract_images(para, doc.part, images_dir, f"p{position}")

        if level > 0 and text:
            # Flush current section before starting a new heading block.
            if current_text_parts or current_images:
                sections.append(
                    DocSection(
                        heading=current_heading,
                        heading_level=current_level,
                        text="\n".join(current_text_parts).strip(),
                        images=current_images,
                        position=position,
                    )
                )
            current_heading = text
            current_level = level
            current_text_parts = []
            current_images = list(para_images)
        else:
            if text:
                current_text_parts.append(text)
            current_images.extend(para_images)

    # Flush the last section.
    if current_text_parts or current_images:
        sections.append(
            DocSection(
                heading=current_heading,
                heading_level=current_level,
                text="\n".join(current_text_parts).strip(),
                images=current_images,
                position=len(doc.paragraphs),
            )
        )

    return sections


def _parse_tables(doc: Document, start_position: int) -> list[DocSection]:
    """Convert tables to text sections so they are also indexed."""
    sections: list[DocSection] = []
    for idx, table in enumerate(doc.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        text = "\n".join(rows).strip()
        if text:
            sections.append(
                DocSection(
                    heading=f"Таблица {idx + 1}",
                    heading_level=0,
                    text=text,
                    images=[],
                    position=start_position + idx,
                )
            )
    return sections


def _heading_level(style_name: str) -> int:
    """Return 1–9 for heading styles, 0 otherwise."""
    name = style_name.lower().strip()
    for level in range(1, 10):
        if name in (f"heading {level}", f"heading{level}",
                    f"заголовок {level}", f"заголовок{level}"):
            return level
    # Fallback: style starts with known prefix.
    if name.startswith(("heading", "заголовок")):
        return 1
    return 0


def _extract_images(para, doc_part, images_dir: str, prefix: str) -> list[str]:
    """
    Extract and save all inline images found in a paragraph element.

    Looks for ``a:blip`` elements which carry the relationship ID pointing to
    the actual image binary stored in the docx package.
    """
    saved: list[str] = []
    blips = para._element.findall(".//" + qn("a:blip"))
    for idx, blip in enumerate(blips):
        r_id = blip.get(qn("r:embed"))
        if not r_id:
            continue
        try:
            rel = doc_part.rels.get(r_id)
            if rel is None or "image" not in rel.reltype:
                continue
            data: bytes = rel.target_part.blob
            ext = rel.target_ref.rsplit(".", 1)[-1].lower() if "." in rel.target_ref else "png"
            filename = f"{prefix}_img{idx}.{ext}"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as fh:
                fh.write(data)
            saved.append(filepath)
            logger.debug("Extracted image → %s", filepath)
        except Exception as exc:
            logger.warning("Could not extract image rId=%s: %s", r_id, exc)
    return saved
