"""
Loads and parses knowledge base files (.docx, .pdf, .xlsx).

Returns a list of DocSection objects, each corresponding to one logical section.
Inline images are extracted and saved to *images_dir* with names that include
a section/paragraph index so they can later be linked back to chunks.
"""
from __future__ import annotations

import html as html_module
import logging
import os
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


@dataclass
class DocSection:
    """One logical section: a heading + all body text until the next heading."""

    heading: str
    heading_level: int        # 0 = no heading, 1 = Heading 1, ...
    text: str                 # concatenated paragraph text
    images: list[str] = field(default_factory=list)  # absolute paths to images
    position: int = 0         # paragraph index of the heading in the document
    source_file: str = ""     # cleaned filename for context


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_knowledge_base_multi(file_paths: list[str], images_dir: str) -> list[DocSection]:
    """Parse multiple files of any supported format and return combined sections."""
    all_sections: list[DocSection] = []
    for path in file_paths:
        try:
            ext = Path(path).suffix.lower()
            if ext == ".docx":
                sections = _load_docx(path, images_dir)
            elif ext == ".pdf":
                sections = _load_pdf(path, images_dir)
            elif ext == ".xlsx":
                sections = _load_xlsx(path, images_dir)
            else:
                logger.debug("Skipping unsupported format: %s", path)
                continue
            all_sections.extend(sections)
            logger.info("Loaded %d sections from '%s'", len(sections), Path(path).name)
        except Exception as e:
            logger.warning("Skipping file %s: %s", path, e)
    return all_sections


# Keep backward-compatible alias
def load_knowledge_base(docx_path: str, images_dir: str) -> list[DocSection]:
    return _load_docx(docx_path, images_dir)


# ---------------------------------------------------------------------------
# DOCX loader
# ---------------------------------------------------------------------------

def _load_docx(docx_path: str, images_dir: str) -> list[DocSection]:
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {docx_path}")
    if path.stat().st_size == 0:
        raise ValueError(f"File is empty: {docx_path}")

    Path(images_dir).mkdir(parents=True, exist_ok=True)
    source_name = _clean_filename(path)

    try:
        doc = Document(str(path))
    except Exception as exc:
        raise ValueError(f"Cannot open '{docx_path}': {exc}") from exc

    sections = _parse_paragraphs(doc, images_dir, source_name)
    table_sections = _parse_tables(doc, start_position=len(doc.paragraphs), source_name=source_name)
    sections.extend(table_sections)

    if not sections:
        raise ValueError(f"No readable content found in '{docx_path}'")

    return sections


# ---------------------------------------------------------------------------
# PDF loader
# ---------------------------------------------------------------------------

def _load_pdf(pdf_path: str, images_dir: str) -> list[DocSection]:
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {pdf_path}")

    import pdfplumber

    Path(images_dir).mkdir(parents=True, exist_ok=True)
    source_name = _clean_filename(path)

    sections: list[DocSection] = []
    text_parts: list[str] = []
    all_images: list[str] = []

    try:
        with pdfplumber.open(str(path)) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                # Use layout extraction for better spacing
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text and page_text.strip():
                    # Fix common PDF extraction issues: remove excessive whitespace
                    cleaned = re.sub(r"[ \t]+", " ", page_text)
                    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
                    text_parts.append(cleaned.strip())

                # Extract images from page
                for img_idx, img in enumerate(page.images or []):
                    try:
                        img_obj = page.crop(
                            (img["x0"], img["top"], img["x1"], img["bottom"])
                        ).to_image(resolution=150)
                        filename = f"pdf_{path.stem}_p{page_idx}_img{img_idx}.png"
                        filepath = os.path.join(images_dir, filename)
                        img_obj.save(filepath)
                        all_images.append(filepath)
                    except Exception:
                        pass
    except Exception as exc:
        raise ValueError(f"Cannot read PDF '{pdf_path}': {exc}") from exc

    full_text = "\n\n".join(text_parts).strip()
    if not full_text:
        raise ValueError(f"No readable text in PDF '{pdf_path}'")

    # Try to extract a title from first non-empty line
    first_line = full_text.split("\n")[0].strip()
    heading = first_line if len(first_line) < 100 else source_name

    sections.append(
        DocSection(
            heading=heading,
            heading_level=0,
            text=full_text,
            images=all_images,
            position=0,
            source_file=source_name,
        )
    )

    return sections


# ---------------------------------------------------------------------------
# XLSX loader
# ---------------------------------------------------------------------------

def _load_xlsx(xlsx_path: str, images_dir: str = "") -> list[DocSection]:
    path = Path(xlsx_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {xlsx_path}")

    from openpyxl import load_workbook

    source_name = _clean_filename(path)

    try:
        wb = load_workbook(str(path), data_only=True)
    except Exception as exc:
        raise ValueError(f"Cannot open XLSX '{xlsx_path}': {exc}") from exc

    if images_dir:
        Path(images_dir).mkdir(parents=True, exist_ok=True)

    sections: list[DocSection] = []

    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [_strip_html(str(c)).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                rows.append(" | ".join(cells))

        text = "\n".join(rows).strip()
        if not text:
            continue

        # Extract embedded images from the sheet
        sheet_images: list[str] = []
        if images_dir:
            sheet_images = _extract_xlsx_images(ws, images_dir, path.stem, sheet_idx)

        # Use sheet name as heading if meaningful, otherwise filename
        heading = sheet_name if sheet_name.lower() not in ("sheet1", "лист1", "лист 1") else source_name

        sections.append(
            DocSection(
                heading=heading,
                heading_level=0,
                text=text,
                images=sheet_images,
                position=sheet_idx,
                source_file=source_name,
            )
        )

    wb.close()

    if not sections:
        raise ValueError(f"No readable content in XLSX '{xlsx_path}'")

    return sections


def _extract_xlsx_images(ws, images_dir: str, file_stem: str, sheet_idx: int) -> list[str]:
    """Extract embedded images from an xlsx worksheet and save to disk."""
    saved: list[str] = []
    try:
        images = ws._images if hasattr(ws, "_images") else []
    except Exception:
        return saved

    for idx, img in enumerate(images):
        try:
            data = img._data()
            if not data:
                continue
            fmt = getattr(img, "format", "png") or "png"
            filename = f"xlsx_{file_stem}_s{sheet_idx}_img{idx}.{fmt}"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as fh:
                fh.write(data)
            saved.append(filepath)
        except Exception as exc:
            logger.debug("Could not extract xlsx image %d: %s", idx, exc)
    if saved:
        logger.debug("Extracted %d images from xlsx sheet %d", len(saved), sheet_idx)
    return saved


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _clean_filename(path: Path) -> str:
    """Convert filename to a human-readable label without extension."""
    name = path.stem
    # Replace underscores with spaces
    name = name.replace("_", " ")
    # Remove version suffixes like (1), (2)
    name = re.sub(r"\s*\(\d+\)\s*$", "", name)
    # Collapse multiple spaces
    name = re.sub(r"\s+", " ", name).strip()
    return name


def _strip_html(text: str) -> str:
    """Remove HTML tags and decode entities from text."""
    if "<" not in text:
        return text
    # Remove tags
    clean = re.sub(r"<[^>]+>", " ", text)
    # Decode HTML entities
    clean = html_module.unescape(clean)
    # Collapse whitespace
    clean = re.sub(r"\s+", " ", clean).strip()
    return clean


def _parse_paragraphs(doc: Document, images_dir: str, source_name: str) -> list[DocSection]:
    sections: list[DocSection] = []

    # Check if ANY paragraph has a heading style
    has_headings = any(
        _heading_level(p.style.name if p.style else "") > 0 and p.text.strip()
        for p in doc.paragraphs
    )

    # Default heading: use filename when no heading styles found
    default_heading = "Общая информация" if has_headings else source_name

    current_heading = default_heading
    current_level = 0
    current_text_parts: list[str] = []
    current_images: list[str] = []

    for position, para in enumerate(doc.paragraphs):
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        level = _heading_level(style_name)
        para_images = _extract_images(para, doc.part, images_dir, f"p{position}")

        if level > 0 and text:
            # Flush current section before starting a new heading block.
            if current_text_parts or current_images:
                sections.append(
                    DocSection(
                        heading=current_heading,
                        heading_level=current_level,
                        text="\n".join(current_text_parts).strip(),
                        images=current_images,
                        position=position,
                        source_file=source_name,
                    )
                )
            current_heading = text
            current_level = level
            current_text_parts = []
            current_images = list(para_images)
        else:
            if text:
                current_text_parts.append(text)
            current_images.extend(para_images)

    # Flush the last section.
    if current_text_parts or current_images:
        sections.append(
            DocSection(
                heading=current_heading,
                heading_level=current_level,
                text="\n".join(current_text_parts).strip(),
                images=current_images,
                position=len(doc.paragraphs),
                source_file=source_name,
            )
        )

    return sections


def _parse_tables(doc: Document, start_position: int, source_name: str) -> list[DocSection]:
    """Convert tables to text sections so they are also indexed."""
    sections: list[DocSection] = []
    for idx, table in enumerate(doc.tables):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                rows.append(" | ".join(cells))
        text = "\n".join(rows).strip()
        if text:
            sections.append(
                DocSection(
                    heading=f"Таблица {idx + 1}",
                    heading_level=0,
                    text=text,
                    images=[],
                    position=start_position + idx,
                    source_file=source_name,
                )
            )
    return sections


def _heading_level(style_name: str) -> int:
    """Return 1-9 for heading styles, 0 otherwise."""
    name = style_name.lower().strip()
    for level in range(1, 10):
        if name in (f"heading {level}", f"heading{level}",
                    f"заголовок {level}", f"заголовок{level}"):
            return level
    if name.startswith(("heading", "заголовок")):
        return 1
    return 0


def _extract_images(para, doc_part, images_dir: str, prefix: str) -> list[str]:
    """Extract and save all inline images found in a paragraph element."""
    saved: list[str] = []
    blips = para._element.findall(".//" + qn("a:blip"))
    for idx, blip in enumerate(blips):
        r_id = blip.get(qn("r:embed"))
        if not r_id:
            continue
        try:
            rel = doc_part.rels.get(r_id)
            if rel is None or "image" not in rel.reltype:
                continue
            data: bytes = rel.target_part.blob
            ext = rel.target_ref.rsplit(".", 1)[-1].lower() if "." in rel.target_ref else "png"
            filename = f"{prefix}_img{idx}.{ext}"
            filepath = os.path.join(images_dir, filename)
            with open(filepath, "wb") as fh:
                fh.write(data)
            saved.append(filepath)
            logger.debug("Extracted image -> %s", filepath)
        except Exception as exc:
            logger.warning("Could not extract image rId=%s: %s", r_id, exc)
    return saved
